import streamlit as st
from PIL import Image, ImageOps, UnidentifiedImageError
import io
import zipfile
import sys
import subprocess
from datetime import datetime
from typing import List, Tuple

# Optional dependencies
try:
    import fitz  # PyMuPDF für PDF → Word
except ImportError:
    fitz = None

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None
    Pt = None

# =========================================================
# WERTFILE 1.2 CLEAN
# Fokus: stabil, clean, keine unnötigen Elemente
# Funktionen:
# 1. Bild → PDF
# 2. PDF → Word für Text-PDFs
# 3. PDF Merge
# 4. PDF Split
# =========================================================

APP_NAME = "Wertfile"
MAX_FILES = 50
MAX_TOTAL_MB = 200
MAX_SINGLE_MB = 25
SUPPORTED_IMAGE_TYPES = ["jpg", "jpeg", "png", "webp"]
SUPPORTED_PDF_TYPES = ["pdf"]

st.set_page_config(
    page_title=f"{APP_NAME} | File Converter",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

if "active_tool" not in st.session_state:
    st.session_state.active_tool = "image_to_pdf"


# ---------- CSS ----------
def inject_css() -> None:
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

        :root {
            --bg: #F7F9FC;
            --card: #FFFFFF;
            --ink: #0F172A;
            --muted: #64748B;
            --line: #E2E8F0;
            --blue: #2563EB;
            --blue-soft: #EFF6FF;
            --purple: #7C3AED;
            --purple-soft: #F5F3FF;
            --green: #10B981;
            --green-soft: #ECFDF5;
            --pink: #DB2777;
            --pink-soft: #FDF2F8;
            --orange: #F97316;
            --orange-soft: #FFF7ED;
            --shadow: 0 16px 44px rgba(15, 23, 42, 0.06);
            --radius: 24px;
        }

        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif !important;
        }

        .stApp {
            background:
                radial-gradient(circle at 8% 10%, rgba(37,99,235,0.08), transparent 26%),
                radial-gradient(circle at 92% 6%, rgba(16,185,129,0.08), transparent 28%),
                linear-gradient(180deg, #F8FAFC 0%, #F2F6FB 100%) !important;
            color: var(--ink);
        }

        header, footer, .stDeployButton, #MainMenu {
            display: none !important;
        }

        .block-container {
            max-width: 1120px;
            padding-top: 34px !important;
            padding-bottom: 46px !important;
        }

        h1, h2, h3, h4, p, span, label, div {
            color: var(--ink);
        }

        p { line-height: 1.55; }

        .wf-topbar {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 16px;
            margin-bottom: 18px;
        }

        .wf-brand {
            display: flex;
            align-items: center;
            gap: 11px;
            font-weight: 900;
            font-size: 24px;
            letter-spacing: -0.8px;
        }

        .wf-logo {
            width: 40px;
            height: 40px;
            border-radius: 14px;
            display: grid;
            place-items: center;
            color: #FFFFFF !important;
            font-weight: 900;
            background: linear-gradient(135deg, #2563EB 0%, #7C3AED 52%, #10B981 100%);
            box-shadow: 0 12px 26px rgba(37,99,235,0.20);
        }

        .wf-version {
            display: inline-flex;
            align-items: center;
            gap: 8px;
            border-radius: 999px;
            padding: 8px 12px;
            background: #FFFFFF;
            border: 1px solid var(--line);
            color: #475569 !important;
            font-size: 13px;
            font-weight: 800;
        }

        .wf-dot {
            width: 8px;
            height: 8px;
            border-radius: 999px;
            background: var(--green);
            box-shadow: 0 0 0 5px rgba(16,185,129,0.12);
        }

        .wf-title-card {
            background: rgba(255,255,255,0.92);
            border: 1px solid var(--line);
            border-radius: 28px;
            padding: 26px;
            box-shadow: var(--shadow);
            margin-bottom: 16px;
        }

        .wf-title-card h1 {
            font-size: clamp(30px, 4vw, 48px);
            line-height: 1.02;
            letter-spacing: -2px;
            margin: 0 0 10px 0;
            font-weight: 900;
        }

        .wf-gradient {
            background: linear-gradient(90deg, #2563EB, #7C3AED, #10B981);
            -webkit-background-clip: text;
            background-clip: text;
            color: transparent !important;
        }

        .wf-title-card p {
            max-width: 780px;
            margin: 0;
            color: var(--muted) !important;
            font-size: 16px;
            font-weight: 550;
        }

        .wf-tool-strip {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 10px;
            margin-bottom: 16px;
        }

        .wf-tool-box {
            background: #FFFFFF;
            border: 1px solid var(--line);
            border-radius: 20px;
            padding: 14px;
            box-shadow: 0 10px 28px rgba(15,23,42,0.04);
            min-height: 86px;
        }

        .wf-tool-box.active-blue {
            border-color: rgba(37,99,235,0.34);
            background: linear-gradient(145deg, #FFFFFF, var(--blue-soft));
        }

        .wf-tool-box.active-purple {
            border-color: rgba(124,58,237,0.34);
            background: linear-gradient(145deg, #FFFFFF, var(--purple-soft));
        }

        .wf-tool-box.active-green {
            border-color: rgba(16,185,129,0.34);
            background: linear-gradient(145deg, #FFFFFF, var(--green-soft));
        }

        .wf-tool-box.active-orange {
            border-color: rgba(249,115,22,0.34);
            background: linear-gradient(145deg, #FFFFFF, var(--orange-soft));
        }

        .wf-tool-box b {
            display: block;
            font-size: 14px;
            margin-bottom: 4px;
        }

        .wf-tool-box span {
            color: var(--muted) !important;
            font-size: 12px;
            font-weight: 650;
        }

        .wf-card {
            background: rgba(255,255,255,0.94);
            border: 1px solid var(--line);
            border-radius: var(--radius);
            padding: 22px;
            box-shadow: var(--shadow);
            margin-bottom: 16px;
        }

        .wf-card h2, .wf-card h3 {
            margin-top: 0;
            letter-spacing: -0.7px;
        }

        .wf-card-subtitle {
            color: var(--muted) !important;
            margin-top: -8px;
            margin-bottom: 16px;
            font-size: 14px;
            font-weight: 550;
        }

        .wf-info-grid {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 10px;
            margin-top: 14px;
        }

        .wf-info-box {
            background: #F8FAFC;
            border: 1px solid var(--line);
            border-radius: 16px;
            padding: 13px;
        }

        .wf-info-box b {
            display: block;
            font-size: 18px;
        }

        .wf-info-box span {
            color: var(--muted) !important;
            font-size: 12px;
            font-weight: 700;
        }

        .wf-file-list {
            display: grid;
            gap: 8px;
            margin-top: 12px;
        }

        .wf-file-row {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            padding: 12px 13px;
            border-radius: 15px;
            background: #F8FAFC;
            border: 1px solid var(--line);
            font-size: 13px;
            font-weight: 700;
        }

        .wf-file-row small {
            color: var(--muted) !important;
            font-weight: 700;
        }

        .wf-preview-empty {
            min-height: 280px;
            border-radius: 20px;
            background: linear-gradient(145deg, #FFFFFF, #F8FAFC);
            border: 1px dashed #CBD5E1;
            display: grid;
            place-items: center;
            text-align: center;
            padding: 24px;
        }

        .wf-preview-empty h3 {
            margin-bottom: 6px;
        }

        .wf-preview-empty p {
            color: var(--muted) !important;
            margin: 0;
            max-width: 340px;
        }

        .wf-test-card {
            border-radius: 20px;
            padding: 18px;
            background: #F8FAFC;
            border: 1px solid var(--line);
        }

        .wf-test-card ul {
            margin-bottom: 0;
            color: #334155;
        }

        .wf-footer {
            text-align: center;
            color: #94A3B8 !important;
            font-size: 12px;
            font-weight: 700;
            margin-top: 16px;
        }

        .stFileUploader section {
            border-radius: 20px !important;
            border: 1.5px dashed #CBD5E1 !important;
            background: #FFFFFF !important;
            padding: 26px !important;
        }

        .stFileUploader section:hover {
            border-color: #2563EB !important;
            background: #F8FAFC !important;
        }

        .stTextInput input,
        .stNumberInput input,
        .stSelectbox [data-baseweb="select"],
        .stSelectbox [data-baseweb="select"] > div,
        .stTextArea textarea {
            min-height: 48px !important;
            border-radius: 14px !important;
            border-color: #CBD5E1 !important;
            background: #FFFFFF !important;
            color: #0F172A !important;
            font-weight: 700 !important;
        }

        .stFileUploader button,
        .stFileUploader [data-testid="baseButton-secondary"],
        .stFileUploader [data-testid="stBaseButton-secondary"] {
            border-radius: 14px !important;
            min-height: 44px !important;
            border: 1px solid rgba(37,99,235,0.24) !important;
            background: linear-gradient(135deg, #EFF6FF 0%, #F5F3FF 100%) !important;
            color: #1D4ED8 !important;
            font-weight: 850 !important;
            box-shadow: 0 8px 20px rgba(37,99,235,0.10) !important;
        }

        .stFileUploader button:hover,
        .stFileUploader [data-testid="baseButton-secondary"]:hover,
        .stFileUploader [data-testid="stBaseButton-secondary"]:hover {
            border-color: rgba(37,99,235,0.40) !important;
            background: linear-gradient(135deg, #DBEAFE 0%, #EDE9FE 100%) !important;
            color: #1E40AF !important;
        }

        .stRadio > label { display: none !important; }

        .stRadio div[role="radiogroup"] {
            display: grid !important;
            grid-template-columns: repeat(4, 1fr) !important;
            gap: 10px !important;
            margin-bottom: 16px !important;
        }

        .stRadio div[role="radiogroup"] label {
            background: #FFFFFF !important;
            border: 1px solid #E2E8F0 !important;
            border-radius: 16px !important;
            padding: 12px 13px !important;
            box-shadow: 0 8px 22px rgba(15,23,42,0.04) !important;
        }

        .stRadio div[role="radiogroup"] label:has(input:checked) {
            background: linear-gradient(135deg, #EFF6FF 0%, #F5F3FF 100%) !important;
            border-color: rgba(37,99,235,0.36) !important;
            box-shadow: 0 12px 30px rgba(37,99,235,0.10) !important;
        }

        .stRadio div[role="radiogroup"] label p {
            color: #0F172A !important;
            font-weight: 850 !important;
            font-size: 13px !important;
        }

        .stRadio input { accent-color: #2563EB !important; }

        .stButton > button,
        .stDownloadButton > button {
            border-radius: 14px !important;
            min-height: 50px !important;
            border: 1px solid rgba(37,99,235,0.22) !important;
            background: linear-gradient(135deg, #EFF6FF 0%, #F5F3FF 100%) !important;
            color: #1D4ED8 !important;
            font-weight: 850 !important;
            box-shadow: 0 10px 24px rgba(37,99,235,0.10) !important;
            transition: transform .2s ease, box-shadow .2s ease, border-color .2s ease !important;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            transform: translateY(-1px);
            border-color: rgba(37,99,235,0.38) !important;
            box-shadow: 0 14px 34px rgba(37,99,235,0.16) !important;
        }

        .stButton > button:disabled {
            background: #F1F5F9 !important;
            color: #94A3B8 !important;
            box-shadow: none !important;
            border-color: #E2E8F0 !important;
        }

        .stAlert { border-radius: 16px !important; }

        [data-testid="stExpander"] {
            border: 1px solid var(--line) !important;
            border-radius: 18px !important;
            background: #FFFFFF !important;
            box-shadow: 0 8px 24px rgba(15,23,42,0.04) !important;
        }

        @media (max-width: 900px) {
            .wf-topbar { align-items: flex-start; flex-direction: column; }
            .wf-tool-strip { grid-template-columns: 1fr 1fr; }
            .stRadio div[role="radiogroup"] { grid-template-columns: 1fr 1fr !important; }
            .wf-info-grid { grid-template-columns: 1fr; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ---------- Auto Installer ----------
def install_missing_packages(packages: List[str]) -> Tuple[bool, str]:
    try:
        command = [sys.executable, "-m", "pip", "install", *packages]
        result = subprocess.run(command, capture_output=True, text=True, timeout=180)
        stdout = result.stdout or ""
        stderr = result.stderr or ""
        output = stdout + chr(10) + stderr
        return result.returncode == 0, output
    except Exception as exc:
        return False, str(exc)


def render_missing_package_box(title: str, packages: List[str]) -> None:
    st.markdown('<div class="wf-card">', unsafe_allow_html=True)
    st.markdown(f"<h2>{title}</h2>", unsafe_allow_html=True)
    st.warning("Für dieses Tool fehlen noch Python-Pakete. Du musst dafür nicht ins Terminal — Wertfile kann versuchen, sie automatisch in derselben Umgebung zu installieren.")
    st.code(f"Installiert wird: {' '.join(packages)}")
    if st.button("Fehlende Pakete automatisch installieren", use_container_width=True):
        with st.spinner("Pakete werden installiert. Das kann kurz dauern..."):
            success, output = install_missing_packages(packages)
            if success:
                st.success("Installation abgeschlossen. Bitte die App einmal stoppen und neu starten.")
                st.code("streamlit run app.py")
            else:
                st.error("Automatische Installation hat nicht geklappt. Hier ist die Fehlermeldung:")
                st.text_area("Fehlerdetails", output, height=220)
    st.markdown('</div>', unsafe_allow_html=True)


# ---------- Helpers ----------
def mb(size_bytes: int) -> float:
    return size_bytes / (1024 * 1024)


class CachedUpload(io.BytesIO):
    def __init__(self, data: bytes, name: str, file_type: str = "application/pdf"):
        super().__init__(data)
        self.name = name
        self.type = file_type
        self.size = len(data)


def cache_files(cache_key: str, uploaded_files) -> None:
    if not uploaded_files:
        return

    files = uploaded_files if isinstance(uploaded_files, list) else [uploaded_files]
    cached = []
    for file in files:
        file.seek(0)
        data = file.read()
        cached.append({
            "name": file.name,
            "type": getattr(file, "type", "application/octet-stream"),
            "data": data,
        })
        file.seek(0)

    st.session_state[cache_key] = cached


def get_cached_files(cache_key: str) -> List[CachedUpload]:
    cached = st.session_state.get(cache_key, [])
    return [CachedUpload(item["data"], item["name"], item.get("type", "application/pdf")) for item in cached]


def clear_cached_files(cache_key: str) -> None:
    if cache_key in st.session_state:
        del st.session_state[cache_key]


def validate_image_files(uploaded_files) -> Tuple[List[str], List[str]]:
    warnings = []
    errors = []
    if not uploaded_files:
        return warnings, errors
    if len(uploaded_files) > MAX_FILES:
        errors.append(f"Bitte maximal {MAX_FILES} Dateien hochladen.")
    total_mb = sum(mb(file.size) for file in uploaded_files)
    if total_mb > MAX_TOTAL_MB:
        errors.append(f"Die Gesamtgröße liegt bei {total_mb:.1f} MB. Erlaubt sind maximal {MAX_TOTAL_MB} MB.")
    for file in uploaded_files:
        file_mb = mb(file.size)
        if file_mb > MAX_SINGLE_MB:
            errors.append(f"{file.name} ist {file_mb:.1f} MB groß. Pro Datei sind maximal {MAX_SINGLE_MB} MB empfohlen/erlaubt.")
        ext = file.name.split(".")[-1].lower() if "." in file.name else ""
        if ext not in SUPPORTED_IMAGE_TYPES:
            errors.append(f"{file.name} hat ein nicht unterstütztes Format.")
    if len(uploaded_files) > 20:
        warnings.append("Viele Dateien können die Verarbeitung verlangsamen. Kleinere Batches sind stabiler.")
    return warnings, errors


def validate_pdf_file(uploaded_file) -> List[str]:
    errors = []
    if not uploaded_file:
        return errors
    file_mb = mb(uploaded_file.size)
    if file_mb > MAX_TOTAL_MB:
        errors.append(f"Die PDF ist {file_mb:.1f} MB groß. Erlaubt sind maximal {MAX_TOTAL_MB} MB.")
    ext = uploaded_file.name.split(".")[-1].lower() if "." in uploaded_file.name else ""
    if ext != "pdf":
        errors.append("Bitte eine PDF-Datei hochladen.")
    return errors


def validate_pdf_files(uploaded_files, min_files: int = 1) -> List[str]:
    errors = []
    if not uploaded_files:
        return errors
    if len(uploaded_files) < min_files:
        errors.append(f"Bitte mindestens {min_files} PDF-Dateien hochladen.")
    if len(uploaded_files) > MAX_FILES:
        errors.append(f"Bitte maximal {MAX_FILES} Dateien hochladen.")
    total_mb = sum(mb(file.size) for file in uploaded_files)
    if total_mb > MAX_TOTAL_MB:
        errors.append(f"Die Gesamtgröße liegt bei {total_mb:.1f} MB. Erlaubt sind maximal {MAX_TOTAL_MB} MB.")
    for file in uploaded_files:
        ext = file.name.split(".")[-1].lower() if "." in file.name else ""
        if ext != "pdf":
            errors.append(f"{file.name} ist keine PDF-Datei.")
    return errors


def open_image_safely(file) -> Image.Image:
    try:
        image = Image.open(file)
        image = ImageOps.exif_transpose(image)
        return image.convert("RGB")
    except UnidentifiedImageError:
        raise ValueError(f"{file.name} konnte nicht als Bild erkannt werden.")
    except Exception as exc:
        raise ValueError(f"{file.name} konnte nicht gelesen werden: {exc}")


def fit_to_page(image: Image.Image, page_format: str, margin: int, background: str = "white") -> Image.Image:
    if page_format == "Originalgröße":
        return image
    page_sizes = {
        "A4 Hochformat": (1240, 1754),
        "A4 Querformat": (1754, 1240),
        "Letter Hochformat": (1275, 1650),
        "Letter Querformat": (1650, 1275),
    }
    canvas_size = page_sizes.get(page_format, page_sizes["A4 Hochformat"])
    canvas = Image.new("RGB", canvas_size, background)
    max_w = max(100, canvas_size[0] - (margin * 2))
    max_h = max(100, canvas_size[1] - (margin * 2))
    img = image.copy()
    img.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
    x = (canvas_size[0] - img.width) // 2
    y = (canvas_size[1] - img.height) // 2
    canvas.paste(img, (x, y))
    return canvas


def create_pdf(uploaded_files, page_format: str, margin_size: str, sort_mode: str, quality: int) -> bytes:
    files = list(uploaded_files)
    if sort_mode == "Dateiname A–Z":
        files = sorted(files, key=lambda f: f.name.lower())
    elif sort_mode == "Dateiname Z–A":
        files = sorted(files, key=lambda f: f.name.lower(), reverse=True)
    margin_map = {"Keine": 0, "Klein": 40, "Normal": 80, "Groß": 130}
    margin = margin_map.get(margin_size, 80)
    images = []
    for file in files:
        file.seek(0)
        image = open_image_safely(file)
        images.append(fit_to_page(image, page_format, margin))
    if not images:
        raise ValueError("Keine gültigen Bilder vorhanden.")
    buffer = io.BytesIO()
    images[0].save(buffer, format="PDF", save_all=True, append_images=images[1:], resolution=150.0, quality=quality, optimize=True)
    buffer.seek(0)
    return buffer.getvalue()


def create_preview_image(uploaded_file, page_format: str, margin_size: str) -> Image.Image:
    uploaded_file.seek(0)
    image = open_image_safely(uploaded_file)
    margin_map = {"Keine": 0, "Klein": 40, "Normal": 80, "Groß": 130}
    return fit_to_page(image, page_format, margin_map.get(margin_size, 80))


def create_uploaded_zip(uploaded_files) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for file in uploaded_files:
            file.seek(0)
            zip_file.writestr(file.name, file.read())
    buffer.seek(0)
    return buffer.getvalue()


def extract_pdf_text(pdf_file) -> Tuple[List[dict], int, int]:
    if fitz is None:
        raise ImportError("PyMuPDF fehlt. Bitte installiere es mit: pip install pymupdf")
    pdf_file.seek(0)
    pdf_bytes = pdf_file.read()
    pages = []
    total_chars = 0
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        page_count = doc.page_count
        for idx, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            total_chars += len(text)
            pages.append({"page": idx, "text": text})
    return pages, page_count, total_chars


def create_word_from_pdf_text(pdf_file, output_style: str, include_page_numbers: bool) -> Tuple[bytes, int, int]:
    if Document is None:
        raise ImportError("python-docx fehlt. Bitte installiere es mit: pip install python-docx")
    pages, page_count, total_chars = extract_pdf_text(pdf_file)
    if total_chars < 20:
        raise ValueError("In dieser PDF wurde kaum Text gefunden. Wahrscheinlich ist es ein Scan/Bild-PDF. Dafür braucht Wertfile später OCR.")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("Wertfile PDF to Word Export", level=1)
    intro = doc.add_paragraph()
    intro.add_run(f"Quelle: {pdf_file.name}\n").bold = True
    intro.add_run(f"Seiten: {page_count}\n")
    intro.add_run(f"Exportiert am: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")
    for page in pages:
        text = page["text"]
        if not text:
            continue
        if include_page_numbers:
            doc.add_heading(f"Seite {page['page']}", level=2)
        if output_style == "Absätze erhalten":
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            if not paragraphs:
                paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)
        else:
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), page_count, total_chars


def get_pdf_page_count(pdf_file) -> int:
    if fitz is None:
        raise ImportError("PyMuPDF fehlt. Bitte installiere es mit: pip install pymupdf")
    pdf_file.seek(0)
    pdf_bytes = pdf_file.read()
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        return doc.page_count


def merge_pdfs(uploaded_files, sort_mode: str) -> Tuple[bytes, int]:
    if fitz is None:
        raise ImportError("PyMuPDF fehlt. Bitte installiere es mit: pip install pymupdf")

    files = list(uploaded_files)
    if sort_mode == "Dateiname A–Z":
        files = sorted(files, key=lambda f: f.name.lower())
    elif sort_mode == "Dateiname Z–A":
        files = sorted(files, key=lambda f: f.name.lower(), reverse=True)

    output_doc = fitz.open()
    total_pages = 0

    for file in files:
        file.seek(0)
        pdf_bytes = file.read()
        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as input_doc:
                if input_doc.needs_pass:
                    raise ValueError(f"{file.name} ist verschlüsselt und kann nicht verarbeitet werden.")
                output_doc.insert_pdf(input_doc)
                total_pages += input_doc.page_count
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(f"{file.name} konnte nicht verarbeitet werden: {exc}")

    if total_pages == 0:
        output_doc.close()
        raise ValueError("Es wurden keine PDF-Seiten gefunden.")

    merged_bytes = output_doc.tobytes(garbage=4, deflate=True)
    output_doc.close()
    return merged_bytes, total_pages


def parse_page_ranges(range_text: str, max_pages: int) -> List[int]:
    cleaned = range_text.replace(" ", "")
    if not cleaned:
        raise ValueError("Bitte einen Seitenbereich eingeben, z. B. 1-3 oder 1,3,5.")
    selected = []
    parts = cleaned.split(",")
    for part in parts:
        if "-" in part:
            start_str, end_str = part.split("-", 1)
            start = int(start_str)
            end = int(end_str)
            if start > end:
                raise ValueError(f"Ungültiger Bereich: {part}")
            selected.extend(range(start, end + 1))
        else:
            selected.append(int(part))
    unique_pages = []
    for page in selected:
        if page < 1 or page > max_pages:
            raise ValueError(f"Seite {page} liegt außerhalb des Dokuments. Erlaubt: 1 bis {max_pages}.")
        if page not in unique_pages:
            unique_pages.append(page)
    return unique_pages


def split_pdf(pdf_file, range_text: str) -> Tuple[bytes, int, int]:
    if fitz is None:
        raise ImportError("PyMuPDF fehlt. Bitte installiere es mit: pip install pymupdf")

    pdf_file.seek(0)
    pdf_bytes = pdf_file.read()

    try:
        input_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"PDF konnte nicht geöffnet werden: {exc}")

    if input_doc.needs_pass:
        input_doc.close()
        raise ValueError("Diese PDF ist verschlüsselt und kann nicht verarbeitet werden.")

    max_pages = input_doc.page_count
    pages = parse_page_ranges(range_text, max_pages)

    output_doc = fitz.open()
    for page_number in pages:
        output_doc.insert_pdf(input_doc, from_page=page_number - 1, to_page=page_number - 1)

    split_bytes = output_doc.tobytes(garbage=4, deflate=True)
    output_doc.close()
    input_doc.close()
    return split_bytes, len(pages), max_pages


# ---------- UI ----------
def render_header() -> None:
    tool_labels = {
        "image_to_pdf": "Bild → PDF",
        "pdf_to_word": "PDF → Word",
        "pdf_merge": "PDF Merge",
        "pdf_split": "PDF Split",
    }
    tool_text = tool_labels.get(st.session_state.active_tool, "Bild → PDF")
    st.markdown(
        f"""
        <div class="wf-topbar">
            <div class="wf-brand">
                <div class="wf-logo">W</div>
                <div>Wertfile.</div>
            </div>
            <div class="wf-version"><span class="wf-dot"></span>1.2 · {tool_text}</div>
        </div>
        <section class="wf-title-card">
            <h1>PDF Workspace.<br><span class="wf-gradient">Clean und zuverlässig.</span></h1>
            <p>Konvertiere Bilder zu PDF, wandle Text-PDFs in Word um, führe PDFs zusammen oder schneide Seitenbereiche aus.</p>
        </section>
        """,
        unsafe_allow_html=True,
    )


def render_tool_selector() -> None:
    active = st.session_state.active_tool
    cls = {
        "image_to_pdf": "active-blue" if active == "image_to_pdf" else "",
        "pdf_to_word": "active-purple" if active == "pdf_to_word" else "",
        "pdf_merge": "active-green" if active == "pdf_merge" else "",
        "pdf_split": "active-orange" if active == "pdf_split" else "",
    }
    st.markdown(
        f"""
        <div class="wf-tool-strip">
            <div class="wf-tool-box {cls['image_to_pdf']}"><b>Bild → PDF</b><span>JPG, PNG oder WEBP als PDF.</span></div>
            <div class="wf-tool-box {cls['pdf_to_word']}"><b>PDF → Word</b><span>Text-PDFs als DOCX.</span></div>
            <div class="wf-tool-box {cls['pdf_merge']}"><b>PDF Merge</b><span>Mehrere PDFs zusammenführen.</span></div>
            <div class="wf-tool-box {cls['pdf_split']}"><b>PDF Split</b><span>Seitenbereich exportieren.</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    labels = ["Bild → PDF", "PDF → Word", "PDF Merge", "PDF Split"]
    tool_map = {"Bild → PDF": "image_to_pdf", "PDF → Word": "pdf_to_word", "PDF Merge": "pdf_merge", "PDF Split": "pdf_split"}
    reverse_map = {v: k for k, v in tool_map.items()}
    selected = st.radio("Tool auswählen", labels, index=labels.index(reverse_map.get(active, "Bild → PDF")), horizontal=True, label_visibility="collapsed")
    st.session_state.active_tool = tool_map[selected]


def render_image_to_pdf_converter() -> None:
    left, right = st.columns([1.1, 0.9], gap="large")
    with left:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>Bild → PDF</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Bilder hochladen, Einstellungen wählen, PDF herunterladen.</p>', unsafe_allow_html=True)
        uploaded_now = st.file_uploader("Bilder hochladen", type=SUPPORTED_IMAGE_TYPES, accept_multiple_files=True, help=f"Unterstützt: {', '.join(SUPPORTED_IMAGE_TYPES).upper()} · Max. {MAX_FILES} Dateien · max. {MAX_TOTAL_MB} MB gesamt.", key="image_to_pdf_upload")
        if uploaded_now:
            cache_files("image_to_pdf_files", uploaded_now)
        uploaded_files = get_cached_files("image_to_pdf_files")
        if uploaded_files:
            st.caption("Bilder bleiben beim Wechsel zwischen Tools erhalten.")
            if st.button("Gemerkte Bilder entfernen", use_container_width=True):
                clear_cached_files("image_to_pdf_files")
                st.rerun()
        warnings, errors = validate_image_files(uploaded_files)
        for warning in warnings:
            st.warning(warning)
        for error in errors:
            st.error(error)
        if uploaded_files:
            total_mb = sum(mb(file.size) for file in uploaded_files)
            st.markdown(f"""<div class="wf-info-grid"><div class="wf-info-box"><b>{len(uploaded_files)}</b><span>Dateien</span></div><div class="wf-info-box"><b>{total_mb:.1f} MB</b><span>Gesamtgröße</span></div><div class="wf-info-box"><b>PDF</b><span>Export</span></div></div>""", unsafe_allow_html=True)
            st.markdown('<div class="wf-file-list">', unsafe_allow_html=True)
            for file in uploaded_files[:10]:
                st.markdown(f'<div class="wf-file-row"><span>{file.name}</span><small>{mb(file.size):.2f} MB</small></div>', unsafe_allow_html=True)
            if len(uploaded_files) > 10:
                st.markdown(f'<div class="wf-file-row"><span>+ {len(uploaded_files) - 10} weitere Dateien</span><small>nicht angezeigt</small></div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<h3>Einstellungen</h3>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            page_format = st.selectbox("Seitenformat", ["A4 Hochformat", "A4 Querformat", "Letter Hochformat", "Letter Querformat", "Originalgröße"], index=0)
        with c2:
            margin_size = st.selectbox("Rand", ["Normal", "Klein", "Keine", "Groß"], index=0)
        c3, c4 = st.columns(2)
        with c3:
            sort_mode = st.selectbox("Reihenfolge", ["Upload-Reihenfolge", "Dateiname A–Z", "Dateiname Z–A"], index=0)
        with c4:
            quality_label = st.selectbox("Qualität", ["Hoch", "Ausgewogen", "Kleinere Datei"], index=1)
        quality = {"Hoch": 95, "Ausgewogen": 85, "Kleinere Datei": 72}[quality_label]
        disabled = not uploaded_files or bool(errors)
        if disabled:
            st.button("PDF erstellen", use_container_width=True, disabled=True)
        else:
            if st.button("PDF erstellen", use_container_width=True):
                with st.spinner("PDF wird erstellt..."):
                    try:
                        pdf_bytes = create_pdf(uploaded_files, page_format, margin_size, sort_mode, quality)
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                        st.success("PDF erfolgreich erstellt.")
                        st.download_button("PDF herunterladen", data=pdf_bytes, file_name=f"wertfile_export_{timestamp}.pdf", mime="application/pdf", use_container_width=True)
                    except Exception as exc:
                        st.error(str(exc))
        if uploaded_files:
            with st.expander("Originaldateien als ZIP herunterladen"):
                zip_bytes = create_uploaded_zip(uploaded_files)
                st.download_button("Originale als ZIP herunterladen", data=zip_bytes, file_name=f"wertfile_originale_{datetime.now().strftime('%Y%m%d_%H%M')}.zip", mime="application/zip", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    with right:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>Vorschau</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Erste Seite nach deinen Einstellungen.</p>', unsafe_allow_html=True)
        if uploaded_files and not errors:
            preview_files = list(uploaded_files)
            if sort_mode == "Dateiname A–Z":
                preview_files = sorted(preview_files, key=lambda f: f.name.lower())
            elif sort_mode == "Dateiname Z–A":
                preview_files = sorted(preview_files, key=lambda f: f.name.lower(), reverse=True)
            try:
                preview_image = create_preview_image(preview_files[0], page_format, margin_size)
                st.image(preview_image, caption=f"Erste Seite: {preview_files[0].name}", use_container_width=True)
            except Exception as exc:
                st.error(str(exc))
            st.markdown(f"""<div class="wf-info-grid"><div class="wf-info-box"><b>{page_format.split()[0]}</b><span>Format</span></div><div class="wf-info-box"><b>{margin_size}</b><span>Rand</span></div><div class="wf-info-box"><b>{quality_label}</b><span>Qualität</span></div></div>""", unsafe_allow_html=True)
        else:
            st.markdown("""<div class="wf-preview-empty"><div><h3>Noch keine Vorschau</h3><p>Lade mindestens ein Bild hoch. Danach erscheint hier die Vorschau.</p></div></div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


def render_pdf_to_word_converter() -> None:
    if fitz is None or Document is None:
        missing = []
        if fitz is None:
            missing.append("pymupdf")
        if Document is None:
            missing.append("python-docx")
        render_missing_package_box("PDF → Word vorbereiten", missing)
        return
    left, right = st.columns([1.05, 0.95], gap="large")
    with left:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>PDF → Word</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Für Text-PDFs. Gescannte PDFs brauchen später OCR.</p>', unsafe_allow_html=True)
        uploaded_now = st.file_uploader("PDF hochladen", type=SUPPORTED_PDF_TYPES, accept_multiple_files=False, help="Funktioniert am besten mit PDFs, deren Text markierbar/kopierbar ist.", key="pdf_to_word_upload")
        if uploaded_now:
            cache_files("pdf_to_word_file", uploaded_now)
        cached_word_pdfs = get_cached_files("pdf_to_word_file")
        pdf_file = cached_word_pdfs[0] if cached_word_pdfs else None
        if pdf_file:
            st.caption(f"Gemerkte PDF: {pdf_file.name}. Sie bleibt beim Wechsel zwischen Tools erhalten.")
            if st.button("Gemerkte Word-PDF entfernen", use_container_width=True):
                clear_cached_files("pdf_to_word_file")
                st.rerun()
        errors = validate_pdf_file(pdf_file)
        for error in errors:
            st.error(error)
        output_style = st.selectbox("Word-Struktur", ["Absätze erhalten", "Zeilenweise exportieren"], index=0)
        include_page_numbers = st.checkbox("Seitenüberschriften einfügen", value=True)
        disabled = not pdf_file or bool(errors) or fitz is None or Document is None
        if disabled:
            st.button("Word-Datei erstellen", use_container_width=True, disabled=True)
        else:
            if st.button("Word-Datei erstellen", use_container_width=True):
                with st.spinner("Word-Datei wird erstellt..."):
                    try:
                        docx_bytes, page_count, total_chars = create_word_from_pdf_text(pdf_file, output_style, include_page_numbers)
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                        clean_name = pdf_file.name.rsplit(".", 1)[0].replace(" ", "_")
                        st.success(f"Word-Datei erstellt. Seiten: {page_count} · Zeichen: {total_chars}")
                        st.download_button("Word-Datei herunterladen", data=docx_bytes, file_name=f"wertfile_{clean_name}_{timestamp}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    except Exception as exc:
                        st.error(str(exc))
        st.markdown('</div>', unsafe_allow_html=True)
    with right:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>PDF-Analyse</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Zeigt, ob Text erkannt wird.</p>', unsafe_allow_html=True)
        if pdf_file and not errors and fitz is not None:
            try:
                pages, page_count, total_chars = extract_pdf_text(pdf_file)
                text_pages = sum(1 for page in pages if page["text"].strip())
                st.markdown(f"""<div class="wf-info-grid"><div class="wf-info-box"><b>{page_count}</b><span>Seiten</span></div><div class="wf-info-box"><b>{text_pages}</b><span>mit Text</span></div><div class="wf-info-box"><b>{total_chars}</b><span>Zeichen</span></div></div>""", unsafe_allow_html=True)
                if total_chars < 20:
                    st.warning("Kaum Text gefunden. Wahrscheinlich Scan-PDF. OCR folgt später.")
                else:
                    st.success("Text gefunden. PDF eignet sich wahrscheinlich für Word-Export.")
                first_text = ""
                for page in pages:
                    if page["text"].strip():
                        first_text = page["text"].strip()[:2200]
                        break
                if first_text:
                    st.text_area("Text-Vorschau", first_text, height=250)
            except Exception as exc:
                st.error(str(exc))
        else:
            st.markdown("""<div class="wf-preview-empty"><div><h3>Noch keine Analyse</h3><p>Lade eine PDF hoch. Danach siehst du Seitenanzahl, Zeichen und Text-Vorschau.</p></div></div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


def render_pdf_merge_converter() -> None:
    if fitz is None:
        render_missing_package_box("PDF Merge vorbereiten", ["pymupdf"])
        return
    left, right = st.columns([1.05, 0.95], gap="large")
    with left:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>PDF Merge</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Mehrere PDF-Dateien zu einer PDF zusammenführen.</p>', unsafe_allow_html=True)
        uploaded_now = st.file_uploader("PDFs hochladen", type=SUPPORTED_PDF_TYPES, accept_multiple_files=True, help="Lade mindestens zwei PDFs hoch.", key="pdf_merge_upload")
        if uploaded_now:
            cache_files("pdf_workspace_files", uploaded_now)
        pdf_files = get_cached_files("pdf_workspace_files")
        if pdf_files:
            st.caption("Dateien bleiben beim Wechsel zwischen PDF Merge und PDF Split erhalten.")
            if st.button("Gemerkte PDFs entfernen", use_container_width=True):
                clear_cached_files("pdf_workspace_files")
                st.rerun()
        errors = validate_pdf_files(pdf_files, min_files=2)
        for error in errors:
            st.error(error)
        sort_mode = st.selectbox("Reihenfolge", ["Upload-Reihenfolge", "Dateiname A–Z", "Dateiname Z–A"], index=0)
        if pdf_files:
            total_mb = sum(mb(file.size) for file in pdf_files)
            st.markdown(f"""<div class="wf-info-grid"><div class="wf-info-box"><b>{len(pdf_files)}</b><span>PDFs</span></div><div class="wf-info-box"><b>{total_mb:.1f} MB</b><span>Gesamtgröße</span></div><div class="wf-info-box"><b>1</b><span>Export-Datei</span></div></div>""", unsafe_allow_html=True)
            display_files = list(pdf_files)
            if sort_mode == "Dateiname A–Z":
                display_files = sorted(display_files, key=lambda f: f.name.lower())
            elif sort_mode == "Dateiname Z–A":
                display_files = sorted(display_files, key=lambda f: f.name.lower(), reverse=True)
            st.markdown('<div class="wf-file-list">', unsafe_allow_html=True)
            for idx, file in enumerate(display_files, start=1):
                st.markdown(f'<div class="wf-file-row"><span>{idx}. {file.name}</span><small>{mb(file.size):.2f} MB</small></div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        disabled = not pdf_files or bool(errors) or fitz is None
        if disabled:
            st.button("PDFs zusammenführen", use_container_width=True, disabled=True)
        else:
            if st.button("PDFs zusammenführen", use_container_width=True):
                with st.spinner("PDFs werden zusammengeführt..."):
                    try:
                        merged_bytes, total_pages = merge_pdfs(pdf_files, sort_mode)
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                        st.success(f"PDF erfolgreich zusammengeführt. Seiten: {total_pages}")
                        st.download_button("Zusammengeführte PDF herunterladen", data=merged_bytes, file_name=f"wertfile_merge_{timestamp}.pdf", mime="application/pdf", use_container_width=True)
                    except Exception as exc:
                        st.error(str(exc))
        st.markdown('</div>', unsafe_allow_html=True)
    with right:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>Hinweis</h2>", unsafe_allow_html=True)
        st.markdown("""<div class="wf-preview-empty"><div><h3>Reihenfolge prüfen</h3><p>Die PDFs werden in der angezeigten Reihenfolge zusammengeführt. Verschlüsselte PDFs können eventuell nicht verarbeitet werden.</p></div></div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


def render_pdf_split_converter() -> None:
    if fitz is None:
        render_missing_package_box("PDF Split vorbereiten", ["pymupdf"])
        return
    left, right = st.columns([1.05, 0.95], gap="large")
    with left:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>PDF Split</h2>", unsafe_allow_html=True)
        st.markdown('<p class="wf-card-subtitle">Bestimmte Seiten aus einer PDF als neue PDF exportieren.</p>', unsafe_allow_html=True)
        uploaded_now = st.file_uploader("PDF hochladen", type=SUPPORTED_PDF_TYPES, accept_multiple_files=False, key="pdf_split_upload")
        if uploaded_now:
            cache_files("pdf_workspace_files", uploaded_now)
        cached_pdfs = get_cached_files("pdf_workspace_files")
        pdf_file = cached_pdfs[0] if cached_pdfs else None
        if pdf_file:
            st.caption(f"Aktive PDF: {pdf_file.name}. Dateien bleiben beim Wechsel zwischen PDF Merge und PDF Split erhalten.")
            if st.button("Gemerkte PDF entfernen", use_container_width=True):
                clear_cached_files("pdf_workspace_files")
                st.rerun()
        errors = validate_pdf_file(pdf_file)
        for error in errors:
            st.error(error)
        page_count = None
        if pdf_file and not errors and fitz is not None:
            try:
                page_count = get_pdf_page_count(pdf_file)
                st.markdown(f"""<div class="wf-info-grid"><div class="wf-info-box"><b>{page_count}</b><span>Seiten</span></div><div class="wf-info-box"><b>{mb(pdf_file.size):.1f} MB</b><span>Dateigröße</span></div><div class="wf-info-box"><b>PDF</b><span>Export</span></div></div>""", unsafe_allow_html=True)
            except Exception as exc:
                st.error(str(exc))
        range_text = st.text_input("Seitenbereich", placeholder="z. B. 1-3 oder 1,3,5")
        disabled = not pdf_file or bool(errors) or not range_text or fitz is None
        if disabled:
            st.button("Seiten exportieren", use_container_width=True, disabled=True)
        else:
            if st.button("Seiten exportieren", use_container_width=True):
                with st.spinner("Neue PDF wird erstellt..."):
                    try:
                        split_bytes, selected_count, max_pages = split_pdf(pdf_file, range_text)
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                        clean_name = pdf_file.name.rsplit(".", 1)[0].replace(" ", "_")
                        st.success(f"Neue PDF erstellt. Exportierte Seiten: {selected_count} von {max_pages}")
                        st.download_button("Ausgeschnittene PDF herunterladen", data=split_bytes, file_name=f"wertfile_split_{clean_name}_{timestamp}.pdf", mime="application/pdf", use_container_width=True)
                    except Exception as exc:
                        st.error(str(exc))
        st.markdown('</div>', unsafe_allow_html=True)
    with right:
        st.markdown('<div class="wf-card">', unsafe_allow_html=True)
        st.markdown("<h2>Beispiele</h2>", unsafe_allow_html=True)
        st.markdown("""<div class="wf-preview-empty"><div><h3>Seitenbereich eingeben</h3><p><b>1-3</b> exportiert Seite 1 bis 3.<br><b>1,3,5</b> exportiert einzelne Seiten.<br><b>1-2,5</b> kombiniert Bereiche und einzelne Seiten.</p></div></div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


def render_test_plan() -> None:
    with st.expander("Testplan"):
        st.markdown("""
        <div class="wf-test-card">
            <b>Bild → PDF:</b>
            <ul><li>1 JPG hochladen → PDF erstellen → Download öffnen</li><li>3 Bilder hochladen → Reihenfolge prüfen → PDF öffnen</li></ul><br>
            <b>PDF → Word:</b>
            <ul><li>Text-PDF hochladen → Analyse prüfen → Word erstellen</li><li>Scan-PDF testen → OCR-Hinweis sollte erscheinen</li></ul><br>
            <b>PDF Merge:</b>
            <ul><li>2 PDFs hochladen → zusammenführen → Datei öffnen</li><li>Reihenfolge A–Z testen</li></ul><br>
            <b>PDF Split:</b>
            <ul><li>PDF hochladen → 1-2 exportieren</li><li>PDF hochladen → 1,3,5 exportieren</li><li>Ungültige Seite testen → Fehlermeldung prüfen</li></ul>
        </div>
        """, unsafe_allow_html=True)


def render_footer() -> None:
    st.markdown(f'<div class="wf-footer">{APP_NAME}. Version 1.2 Clean · Session-basierte Verarbeitung.</div>', unsafe_allow_html=True)


# ---------- Render ----------
inject_css()
render_header()
render_tool_selector()

if st.session_state.active_tool == "image_to_pdf":
    render_image_to_pdf_converter()
elif st.session_state.active_tool == "pdf_to_word":
    render_pdf_to_word_converter()
elif st.session_state.active_tool == "pdf_merge":
    render_pdf_merge_converter()
else:
    render_pdf_split_converter()

render_test_plan()
render_footer()
